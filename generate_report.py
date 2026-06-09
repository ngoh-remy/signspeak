import os
from docx import Document
from fpdf import FPDF

# Report Content Data Structure
report_data = [
    ("CHAPTER I: GENERAL INTRODUCTION", "h1"),
    ("1.1 Introduction", "h2"),
    ("Sign language is the primary mode of communication for the deaf and hard-of-hearing community. It relies on a complex combination of hand shapes, facial expressions, and body movements. However, a significant communication barrier exists between sign language users and the broader hearing population, who typically do not understand these visual languages. As the world becomes increasingly digital, integrating sign language recognition into modern technology has become a critical step toward accessibility and inclusivity. Recent advancements in artificial intelligence and computer vision now provide the foundation to build systems capable of translating gestures into spoken and written words in real-time.", "p"),
    ("1.2 Background of the Study", "h2"),
    ("Historically, gesture recognition relied on expensive hardware such as sensor-equipped gloves or specialized depth-sensing cameras (e.g., Microsoft Kinect). While effective, these systems were bulky, expensive, and largely restricted to laboratory environments. The advent of deep learning and robust computer vision frameworks has shifted the paradigm toward software-only solutions. By leveraging deep neural networks, developers can now analyze standard 2D webcam feeds to extract 3D skeletal data. Systems have evolved from simple static image classification using Convolutional Neural Networks (CNNs) to dynamic motion tracking using temporal architectures like Long Short-Term Memory (LSTM) networks.", "p"),
    ("1.3 Statement of the Problem", "h2"),
    ("Despite technological advancements, a persistent communication gap isolates the deaf community from critical services, education, and social interaction. Many existing computer vision models for gesture recognition (such as MobileNet) classify individual, static frames. This static approach fundamentally fails when applied to sign language, which is inherently temporal and motion-based. Two different signs may share the exact same hand shape but differ entirely in their trajectory. There is a pressing need for a lightweight, web-accessible system that accurately models the temporal sequences of gestures without requiring users to purchase specialized hardware.", "p"),
    ("1.4 Objectives of the Study", "h2"),
    ("1.4.1 General Objective", "h3"),
    ("The general objective of this study is to design, develop, and deploy a real-time, web-based artificial intelligence system (SignSpeak) capable of translating dynamic sign language sequences into natural language text and speech using a standard webcam.", "p"),
    ("1.4.2 Specific Objectives", "h3"),
    ("- To collect and preprocess a custom dataset of temporal gesture sequences.\n- To extract robust 3D skeletal keypoints (Face, Left Hand, Right Hand) using the MediaPipe Holistic framework.\n- To design and train an LSTM neural network architecture capable of understanding temporal motion patterns.\n- To integrate the trained AI model into a full-stack web application (React frontend, Flask backend) for real-time inference.", "p"),
    ("1.5 Significance of the Study", "h2"),
    ("This study directly promotes accessibility and inclusivity. By providing a scalable, hardware-independent translation tool, it empowers the deaf community to communicate seamlessly in environments such as hospitals, schools, and workplaces. Furthermore, it demonstrates the viability of executing complex spatiotemporal neural networks within standard consumer web browsers.", "p"),
    ("1.6 Scope of the Study", "h2"),
    ("The scope of this project is currently limited to the recognition of 50 core dynamic signs. The system is designed for single-user interaction facing a standard web camera under reasonable lighting conditions. It translates recognized signs into English and French text, accompanied by text-to-speech output.", "p"),
    ("1.7 Definition of Working Terms", "h2"),
    ("- **ASL:** American Sign Language.\n- **MediaPipe:** An open-source framework by Google for cross-platform, customizable machine learning solutions for live and streaming media.\n- **LSTM:** Long Short-Term Memory, a type of recurrent neural network capable of learning order dependence in sequence prediction problems.\n- **Keypoints:** Specific coordinate landmarks detected on a human body (e.g., fingertips, shoulders).\n- **WebRTC:** Web Real-Time Communication, an open framework enabling real-time video streaming in the browser.", "p"),
    ("1.8 Organisation of the Study", "h2"),
    ("This report is structured into several chapters. Chapter I introduces the study, its background, problem statement, and objectives. Chapter II provides a comprehensive review of existing literature and systems, culminating in the proposed solution. Subsequent chapters will detail the system methodology, implementation, and final conclusions.", "p"),
    ("CHAPTER II: LITERATURE REVIEW", "h1"),
    ("2.1 Introduction", "h2"),
    ("This chapter reviews the evolution of gesture recognition systems. It examines the transition from hardware-heavy approaches to modern deep learning architectures, analyzing the strengths and limitations of various methodologies in tackling the complexities of sign language translation.", "p"),
    ("2.2 Review of Existing Systems", "h2"),
    ("2.2.1 Sensor-based Glove Systems", "h3"),
    ("Early systems utilized electromechanical gloves equipped with flex sensors and accelerometers to measure finger bending and hand orientation. While highly accurate, they were expensive and restricted natural hand movement.", "p"),
    ("2.2.2 Depth Cameras and Microsoft Kinect", "h3"),
    ("The introduction of RGB-D cameras like the Kinect allowed for non-intrusive 3D tracking. Researchers successfully used depth maps to segment hands from complex backgrounds, though the hardware remained cumbersome for everyday mobile use.", "p"),
    ("2.2.3 Hidden Markov Models (HMM)", "h3"),
    ("Before the dominance of deep learning, HMMs were the standard for modeling the temporal aspects of gestures. They relied heavily on handcrafted features and struggled to scale with large vocabularies.", "p"),
    ("2.2.4 Convolutional Neural Networks (CNN) for Static Signs", "h3"),
    ("CNNs revolutionized image classification. Many studies applied CNNs to sign language by treating gestures as static images (e.g., fingerspelling alphabets). However, CNNs suffer from \"motion blindness\" when dealing with dynamic signs.", "p"),
    ("2.2.5 3D Convolutional Neural Networks (3D-CNN)", "h3"),
    ("To address motion, researchers developed 3D-CNNs that treat video sequences as 3D volumes. While effective at capturing spatiotemporal features, they are highly computationally expensive and difficult to run in real-time on standard devices.", "p"),
    ("2.2.6 Lightweight Architectures (MobileNet)", "h3"),
    ("MobileNet and similar architectures brought deep learning to mobile devices by reducing parameter counts. Some systems attempt to recognize signs by applying majority voting across multiple MobileNet frame predictions, though this still fails to truly model temporal trajectories.", "p"),
    ("2.2.7 Recurrent Neural Networks (RNN) and LSTMs", "h3"),
    ("LSTMs specifically address the vanishing gradient problem in traditional RNNs, making them ideal for sequence data. By feeding coordinate data frame-by-frame into an LSTM, systems can learn the direction and speed of a gesture over time.", "p"),
    ("2.2.8 MediaPipe Hand Tracking and Holistic Models", "h3"),
    ("Google's MediaPipe provides highly optimized, real-time 3D landmark extraction from 2D video. The Holistic model simultaneously tracks 543 landmarks across the face, hands, and pose, providing an incredibly rich feature set for gesture analysis.", "p"),
    ("2.2.9 Benchmark Datasets (WLASL)", "h3"),
    ("The Word-Level American Sign Language (WLASL) dataset is a massive video dataset containing over 2,000 signs. It serves as a benchmark for modern AI models, though smaller, custom datasets are often required for specialized applications.", "p"),
    ("2.2.10 Web-based AI Inference Architectures", "h3"),
    ("Modern implementations favor client-server architectures where a React frontend captures video via WebRTC, transmits frames (or extracted landmarks) via websockets or REST APIs, and a Python Flask/FastAPI backend executes the neural network inference.", "p"),
    ("2.3 Proposed Solution", "h2"),
    ("Based on the literature, the proposed 'SignSpeak' system abandons static CNN image classification in favor of a dynamic sequence-modeling approach. The system leverages the MediaPipe Holistic framework to extract 258 vital 3D skeletal coordinates (Face, Left Hand, Right Hand) per frame. These coordinates are spatially zero-centered to ensure invariance to the user's position on screen. A sequence of 30 frames is then fed into a Long Short-Term Memory (LSTM) neural network, which accurately deciphers the temporal motion of the sign. This entire AI pipeline is exposed via a Flask API and consumed by a responsive React frontend, delivering a seamless, real-time translation experience directly in the browser.", "p")
]

# Generate Word Document
doc = Document()

for text, element_type in report_data:
    if element_type == "h1":
        doc.add_page_break() if len(doc.paragraphs) > 1 else None
        doc.add_heading(text, level=1)
    elif element_type == "h2":
        doc.add_heading(text, level=2)
    elif element_type == "h3":
        doc.add_heading(text, level=3)
    elif element_type == "p":
        doc.add_paragraph(text)

doc_path = os.path.join(os.path.dirname(__file__), "SignSpeak_Report.docx")
doc.save(doc_path)
print(f"Generated DOCX at: {doc_path}")

# Generate PDF Document
class PDF(FPDF):
    def header(self):
        self.set_font("Arial", "B", 12)
        self.cell(0, 10, "SignSpeak: Project Report", 0, 1, "C")
        self.ln(5)

    def footer(self):
        self.set_y(-15)
        self.set_font("Arial", "I", 8)
        self.cell(0, 10, f"Page {self.page_no()}", 0, 0, "C")

pdf = PDF()
pdf.set_auto_page_break(auto=True, margin=15)
pdf.add_page()

for text, element_type in report_data:
    if element_type == "h1":
        pdf.set_font("Arial", "B", 16)
        if pdf.get_y() > 30: # Don't add page break if it's the very beginning
            pdf.add_page()
        pdf.multi_cell(0, 10, text)
        pdf.ln(5)
    elif element_type == "h2":
        pdf.set_font("Arial", "B", 14)
        pdf.ln(4)
        pdf.multi_cell(0, 10, text)
        pdf.ln(2)
    elif element_type == "h3":
        pdf.set_font("Arial", "B", 12)
        pdf.ln(2)
        pdf.multi_cell(0, 8, text)
        pdf.ln(1)
    elif element_type == "p":
        pdf.set_font("Arial", "", 11)
        # Encode strictly to avoid Latin-1 issues in fpdf
        clean_text = text.replace('\n', ' ').replace('"', "'")
        pdf.multi_cell(0, 6, clean_text)
        pdf.ln(4)

pdf_path = os.path.join(os.path.dirname(__file__), "SignSpeak_Report.pdf")
pdf.output(pdf_path)
print(f"Generated PDF at: {pdf_path}")
